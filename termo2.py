import io
from datetime import date
import streamlit as st
from docx import Document

# ----------------------------------
# Configurações gerais do app
# ----------------------------------
st.set_page_config(
    page_title="Gerador de Termo de Referência — Lei 14.133/2021 (Brasnorte-MT)",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ----------------------------------
# Funções utilitárias
# ----------------------------------
def lista_nao_vazia(valor: str) -> bool:
    return bool(valor and str(valor).strip())

def monta_bloco(label: str, valor: str) -> str:
    """Renderiza um parágrafo só se houver conteúdo."""
    return f"- **{label}:** {valor.strip()}\n" if lista_nao_vazia(valor) else ""

def to_docx(md_text: str) -> bytes:
    """
    Gera um DOCX simples a partir do texto Markdown.
    (Implementação básica: insere como parágrafos; para estilização avançada, parsear o Markdown.)
    """
    doc = Document()
    for line in md_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def formatar_secretarias(lista_sel):
    if not lista_sel:
        return ""
    if len(lista_sel) == 1:
        return lista_sel[0]
    return "; ".join(lista_sel)

def gerar_tr(
    objeto: str,
    secretarias: list[str],
    vigencia_meses: int,
    incluir_opcao_hibrida: bool,
    kpis_padrao: bool,
    municipio: str = "Brasnorte-MT",
) -> str:
    """Gera o TR completo em Markdown com os campos simplificados."""
    hoje = date.today().strftime("%d/%m/%Y")
    objeto_md = f"**{objeto.strip()}**" if lista_nao_vazia(objeto) else "**[INSERIR OBJETO SOLICITADO]**"
    secretarias_md = formatar_secretarias(secretarias)

    md = []
    md.append(f"# TERMO DE REFERÊNCIA — Lei nº 14.133/2021\n")
    md.append(f"**Município:** {municipio}  \n**Data:** {hoje}\n")
    md.append("---\n")

    # 1. Das condições gerais
    md.append("## 1. DAS CONDIÇÕES GERAIS DA CONTRATAÇÃO\n")
    md.append(
        f"1.1 O presente Termo de Referência tem por objeto {objeto_md}, em conformidade com as especificações de descrição e quantidade detalhadamente elencadas neste documento, amparada pelas disposições legais vigentes que regulam tal procedimento, visando atender as necessidades da Prefeitura Municipal de Brasnorte-MT e de suas Secretarias Municipais.\n"
    )
    md.append(f"1.2 O objeto desta contratação não se enquadra como sendo de bem de luxo, conforme Decreto Municipal nº 03/2024.\n")
    md.append(
        f"1.3 O prazo de vigência da contratação será de 12 meses, contados da data de assinatura da ARP (Ata Registro de Preço) ou do Contrato conforme celebrado, na forma do artigo 105 da Lei n° 14.133/2021, podendo o mesmo ser prorrogado a critério da Administração Pública.\n"
    )
    
    md.append(f"1.4 O custo estimado total da contratação é de R$ 00.000,00 (descrever o valor em reais) conforme custos unitários apostos na tabela acima, conforme pesquisa de preço nos termos do Decreto Municipal n° 05/2024.\n"
    )


    # Quadro‑resumo (simplificado)
    md.append("\n**Quadro-resumo do objeto:**\n")
    md.append(monta_bloco("Objeto detalhado", objeto))
    if lista_nao_vazia(secretarias_md):
        md.append(monta_bloco("Unidade(s) demandante(s)", secretarias_md))
    md.append(monta_bloco("Prazo de vigência (meses)", str(vigencia_meses)))
    md.append("\n")

    # 2. Necessidade e fundamentação
    md.append("## 2. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO E FUNDAMENTAÇÃO LEGAL\n")
    import openai

# Insira sua chave de API da OpenAI aqui
openai.api_key = "sua-chave-api-aqui"

def gerar_justificativa(objeto):
    prompt = f"""
Você é um especialista em licitações públicas, com mais de 10 anos de experiência, atuando como responsável técnico pela elaboração e instrução de processos de contratação na Prefeitura Municipal de Brasnorte-MT. Considerando a Lei nº 14.133/2021 e demais normativos aplicáveis, elabore a seção "Justificativa e da Necessidade da Contratação" para compor o Termo de Referência, cujo objeto é: "{objeto}".

O texto deve:
- Apresentar a descrição clara da demanda pública, alinhada às atividades administrativas e operacionais do município e suas secretarias.
- Explicar por que a contratação é necessária, com base em evidências práticas e na ausência de estrutura própria da Administração (quando aplicável).
- Demonstrar que o objeto é essencial à continuidade dos serviços públicos ou à implementação de políticas públicas locais.
- Fundamentar-se nas especificações e quantidades descritas no Termo de Referência.
- Assegurar que a demanda está alinhada ao Plano Anual de Contratações (PAC), e que há previsão orçamentária compatível.
- Apontar a impossibilidade de execução direta pela Administração, se for o caso.
- Incluir os fundamentos legais pertinentes, especialmente o artigo 6º, inciso XXIII, alínea ‘b’, da Lei nº 14.133/2021.
- Ser redigida em linguagem técnica, objetiva e juridicamente fundamentada, voltada à instrução de um processo administrativo de contratação pública.
"""

    resposta = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Você é um redator técnico especialista em licitações públicas."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.4,
        max_tokens=1000
    )

    justificativa = resposta['choices'][0]['message']['content']
    return justificativa


if __name__ == "__main__":
    objeto_input = input("Informe o objeto da contratação (ex: Locação de veículos utilitários com motorista): ")
    texto_gerado = gerar_justificativa(objeto_input)
    print("\n--- JUSTIFICATIVA E DA NECESSIDADE DA CONTRATAÇÃO ---\n")
    print(texto_gerado)

    
    
    
    md.append("- **Contexto e problema a resolver:** descreve por que o objeto é necessário, quem será atendido e quais resultados públicos se pretende alcançar.\n")
    md.append("- **Consequências da não contratação:** riscos operacionais, legais, orçamentários e de continuidade do serviço público.\n")
    md.append("- **Alinhamento ao planejamento:** vinculação a PPA/LDO/LOA e planos setoriais, quando aplicável.\n")
    md.append("- **Justificativa técnica e vantajosidade:** adequação do objeto em desempenho, qualidade, custo total do ciclo de vida e economicidade.\n")
    md.append("- **Fundamentação legal sucinta:** dispositivos pertinentes da Lei nº 14.133/2021 (ex.: art. 6º, art. 40 e, quando cabível, art. 92) e Decreto Municipal nº 09/2024, sem transcrições.\n")

    # 3. Solução, opções e ciclo de vida
    md.append("## 3. DESCRIÇÃO DA SOLUÇÃO COMO UM TODO (CICLO DE VIDA) E ESPECIFICAÇÃO DOS SERVIÇOS\n")
    md.append("**Resumo da necessidade (síntese):** [apresentar em 3–5 linhas].\n\n")
    md.append("**Opções de solução:**\n")
    md.append("- **Opção A — Execução própria pela Prefeitura:** recursos humanos, infraestrutura e competências exigidas; "
              "limitações (ex.: ausência de suporte técnico/equipe qualificada, custos de capacitação, riscos de continuidade) e inviabilidade prática/econômica.\n")
    md.append("- **Opção B — Contratação/Aquisição do objeto:** atendimento pelo mercado, níveis de serviço, prazos, garantias, manutenção/assistência técnica.\n")
    if incluir_opcao_hibrida:
        md.append("- **Opção C — Híbrida/Colaborativa:** parte interna + terceirização de etapas específicas, com avaliação de prós e contras.\n")
    md.append("\n**Conclusão – Solução escolhida:** justificar a alternativa mais vantajosa ao interesse público (eficiência, economicidade e qualidade).\n")
    md.append("**Ciclo de Vida do Objeto:** aquisição/implantação → operação → manutenção/assistência → atualizações/treinamentos → desmobilização/descartes, "
              "incluindo sustentabilidade, garantia e suporte pós‑venda.\n")
    md.append("**Especificação técnica:** características mínimas, desempenho esperado, normas aplicáveis (ABNT/INMETRO/ANVISA/ANEEL etc.), "
              "padrões de qualidade, prazos de atendimento, SLAs e evidências de conformidade.\n")

    # 4. Requisitos
    md.append("## 4. REQUISITOS DA CONTRATAÇÃO\n")
    md.append("Liste requisitos **objetivos e verificáveis** (adapte ao objeto):\n")
    md.append("1. Conformidade técnica com as especificações e normas indicadas.\n")
    md.append("2. Qualificação técnica mínima (atestados, equipes, certificações quando cabíveis).\n")
    md.append("3. Prazos de entrega/execução (SLA, janelas de atendimento, tempo de resposta).\n")
    md.append("4. Garantia (prazo, cobertura, substituição/recall quando aplicável).\n")
    md.append("5. Assistência técnica/manutenção (preventiva e corretiva, tempos de restauração).\n")
    md.append("6. Treinamento/capacitação de usuários/servidores, com material didático.\n")
    md.append("7. Documentação técnica (manuais, catálogos, ART/RRT quando exigível).\n")
    md.append("8. Segurança e conformidade regulatória (saúde, meio ambiente, LGPD quando pertinente).\n")
    md.append("9. Logística e entrega (locais, horários, acondicionamento, rastreabilidade).\n")
    md.append("10. Medição e aceitação (procedimentos, evidências, formulários).\n")
    md.append("11. Sustentabilidade (eficiência energética, redução de resíduos, destinação final).\n")
    md.append("12. Penalidades e garantias contratuais alinhadas à Lei nº 14.133/2021.\n")

    # 5. Execução contratual
    md.append("## 5. MODELO DE EXECUÇÃO CONTRATUAL\n")
    md.append("- **5.1** Execução fiel pelas partes, conforme cláusulas e Lei nº 14.133/2021; responsabilidade por inexecução total ou parcial.\n")
    md.append("- **5.2** Execução conforme este TR, observando Edital e Instrumento Contratual após assinatura.\n")
    md.append("- **5.3** Solicitação do objeto **de forma parcelada**, mediante **OS** e **NE**.\n")
    md.append("- **5.4** Comprovação por **Nota Fiscal** da contratada, **ateste** por servidor competente, com **relatório circunstanciado** (ex.: livro de ponto, comprovantes de entrega/serviços).\n")
    md.append("- **5.5** Responsabilidade integral da contratada pelos ônus de execução.\n")
    md.append("- **5.6** Observância da **NAD (Nota de Autorização de Despesas)**.\n")
    md.append("- **5.7** Comunicações formais **por escrito** (admitido meio eletrônico quando aplicável).\n")
    md.append("- **5.8** Prestação **sob demanda** mediante OS/documento equivalente, com **prazos e quantidades** definidos.\n")
    md.append("- **5.9** Plano de mobilização/desmobilização e cronograma físico‑financeiro (quando aplicável).\n")
    md.append("- **5.10** Gestão e fiscalização contratual (gestor e fiscais; rotinas de reunião e reporte).\n")
    md.append("- **5.11** Confidencialidade, proteção de dados e propriedade intelectual (quando pertinente).\n")
    md.append("- **5.12** Subcontratação e equipe mínima (critérios e limites, quando admitido).\n")
    md.append("- **5.13** Reposição de bens/partes e prazos de correção de não conformidades.\n")
    md.append("- **5.14** Indicadores de desempenho vinculados à medição/aceite e sanções.\n")

    # 6. Critérios de medição
    md.append("## 6. CRITÉRIOS DE MEDIÇÃO\n")
    md.append("- **6.1 Itens e unidades de medida:**\n")
    md.append("  | Item | Descrição | Unidade | Qtde medida no período | Qtde acumulada | Saldo |\n")
    md.append("  |---|---|---|---:|---:|---:|\n")
    md.append("  | 1 | [Descrever] | [un/h/m²/mês] | 0 | 0 | 0 |\n")
    md.append("- **6.2 Evidências de execução:** relatórios, checklists assinados, registros fotográficos, logs/sistemas, canhotos de entrega, certificados de treinamento.\n")
    md.append("- **6.3 Critérios de aceite:** padrões técnicos, tolerâncias e desempenho; procedimento de inspeção (amostragem, testes, prazos para correção).\n")
    md.append("- **6.4 Indicadores de desempenho (SLA/KPI):**\n")
    if kpis_padrao:
        md.append("  - **Disponibilidade (%):** (Horas disponíveis ÷ Horas previstas) × 100.\n")
        md.append("  - **Tempo de resposta (h):** tempo entre abertura e primeiro atendimento.\n")
        md.append("  - **Tempo de solução (h):** tempo entre abertura e solução.\n")
        md.append("  - **Taxa de retrabalho (%):** (Ocorrências retrabalhadas ÷ Total de ocorrências) × 100.\n")
        md.append("  - **Conformidade amostral (%):** (Itens conformes ÷ Itens amostrados) × 100.\n")
        md.append("  - **Pontualidade em entregas (%):** (Entregas pontuais ÷ Entregas totais) × 100.\n")
    else:
        md.append("  - [Definir de 3 a 6 indicadores mensuráveis coerentes com o objeto]\n")
    md.append("- **6.5 Fórmulas de cálculo:** explicitar fórmulas dos indicadores adotados.\n")
    md.append("- **6.6 Periodicidade da medição:** [semanal/mensal/por OS/por marco].\n")
    md.append("- **6.7 Glosas e penalidades:** condições e procedimentos para glosa/desconto, reconvocação, reexecução e penalidades contratuais (sem transcrições legais).\n")
    md.append("- **6.8 Aceite final:** condições para aceite definitivo, termo de recebimento e encerramento.\n")

    md.append("\n---\n")
    md.append("_Observação: este documento deve ser ajustado ao objeto específico, convertendo requisitos em métricas mensuráveis (números, tolerâncias, prazos e padrões)._")

    return "".join(md)

# ----------------------------------
# Lista fixa de Secretarias (multiselect)
# ----------------------------------
SECRETARIAS_PADRAO = [
    "Gabinete Municipal",
    "Secretaria Municipal de Administração",
    "Secretaria Municipal de Assistencia Social",
    "Secretaria Municipal de Assuntos Indigenas",
    "Secretaria Municipal de Des. Agrario e Meio Ambiente",
    "Secretaria Municipal de Educação",
    "Secretaria Municipal de Esportes",
    "Secretaria Municipal de Finanças",
    "Secretaria Municipal de Infraestrutura",
    "Secretaria Municipal de Planejamento, Turismo e Cultura",
    "Secretaria Municipal de Saúde",
]

# ----------------------------------
# UI
# ----------------------------------
st.title("Gerador de Termo de Referência — Lei 14.133/2021 (Pref. de Brasnorte-MT)")
st.caption("Agora com seleção de Secretarias e vigência em meses, conforme solicitado.")

with st.sidebar:
    st.header("Dados de entrada")
    objeto = st.text_input("Objeto detalhado (obrigatório)", placeholder="Ex.: Contratação de serviços de manutenção predial preventiva e corretiva")

    secretarias_sel = st.multiselect(
        "Unidade(s) demandante(s) — selecione as Secretarias participantes do certame",
        options=SECRETARIAS_PADRAO,
        default=[],
        help="Você pode selecionar uma ou várias Secretarias."
    )

    vigencia_meses = st.number_input(
        "Prazo de vigência (meses)",
        min_value=1,
        max_value=60,
        value=12,
        step=1,
        help="Informe apenas o período de vigência do contrato em meses."
    )

    st.markdown("---")
    incluir_opcao_hibrida = st.checkbox("Incluir Opção C (modelo híbrido/colaborativo)", value=True)
    kpis_padrao = st.checkbox("Incluir KPIs/SLAs padrão sugeridos", value=True)
    st.markdown("---")
    gerar = st.button("Gerar Termo de Referência", type="primary", use_container_width=True)

if gerar:
    if not lista_nao_vazia(objeto):
        st.error("Informe o **Objeto detalhado** para gerar o TR.")
        st.stop()

    resultado = gerar_tr(
        objeto=objeto,
        secretarias=secretarias_sel,
        vigencia_meses=int(vigencia_meses),
        incluir_opcao_hibrida=incluir_opcao_hibrida,
        kpis_padrao=kpis_padrao,
        municipio="Brasnorte-MT",
    )

    st.success("TR gerado com sucesso! Revise e ajuste os pontos específicos do objeto antes de publicar.")
    st.download_button(
        label="Baixar em Markdown (.md)",
        data=resultado.encode("utf-8"),
        file_name="TR_Lei_14133_Brasnorte.md",
        mime="text/markdown",
        use_container_width=True,
    )

    docx_bytes = to_docx(resultado)
    st.download_button(
        label="Baixar em Word (.docx)",
        data=docx_bytes,
        file_name="TR_Lei_14133_Brasnorte.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.markdown("### Pré-visualização")
    st.markdown(resultado)
else:
    st.info("Preencha os campos no painel lateral e clique em **Gerar Termo de Referência**.")

st.markdown("---")
st.caption("© Prefeitura Municipal de Brasnorte-MT — Modelo orientado pela Lei nº 14.133/2021. Ajuste conforme o objeto específico e as diretrizes internas.")
